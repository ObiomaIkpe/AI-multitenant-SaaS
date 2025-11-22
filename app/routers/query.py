from fastapi import APIRouter, Depends, UploadFile, File, Form, HTTPException, status
from typing import Optional
from sqlalchemy.orm import Session
from sqlalchemy import select, func
from sqlalchemy.ext.asyncio import AsyncSession
import uuid

from llama_index.core import Document as LlamaDocument, VectorStoreIndex
from llama_index.core.settings import Settings
from llama_index.llms.gemini import Gemini
from llama_index.core.postprocessor import SentenceTransformerRerank
from llama_index.core.query_engine import RetrieverQueryEngine
from llama_index.core.node_parser import SentenceSplitter
from llama_index.vector_stores.qdrant import QdrantVectorStore
from llama_index.core.vector_stores.types import MetadataFilters, ExactMatchFilter

from app.dependencies.common import get_qdrant_client_dependency, get_async_qdrant_client_dependency
from app.core.security import get_current_tenant_id
from app.database import get_db
from app.core.config import settings
from app.models.document import Document, Tenant, User, QueryLog
from app.models.enums import DocumentStatus
from app.schemas.document_schema import (
    DocumentUploadRequest,
    DocumentUploadResponse,
    DocumentResponse,
    DocumentListResponse, 
    DocumentUpdateRequest
)
from datetime import datetime

from app.schemas.query_schema import QueryRequest, QueryResponse
from app.utils.filters import build_metadata_filters, get_applied_filters_summary

router = APIRouter()

@router.post("/upload", response_model=DocumentUploadResponse)
async def upload_document(
    file: UploadFile = File(...),
    category: Optional[str] = Form(None),
    tags: Optional[str] = Form(None),
    description: Optional[str] = Form(None),
    tenant_id: str = Depends(get_current_tenant_id),
    db: AsyncSession = Depends(get_db),
    qdrant_client = Depends(get_qdrant_client_dependency),
    async_qdrant_client = Depends(get_async_qdrant_client_dependency)
):
    """Upload and index a document with metadata for filtering. Supports: PDF, DOCX, TXT files"""
    
    # 1. CHECK TENANT LIMITS
    result = await db.execute(select(Tenant).where(Tenant.id == tenant_id))
    tenant = result.scalar_one_or_none()
    
    if not tenant:
        raise HTTPException(status_code=404, detail="Tenant not found")
    
    result = await db.execute(select(Document).where(Document.tenant_id == tenant_id))
    doc_count = len(result.scalars().all())
    
    if doc_count >= tenant.max_documents:
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail=f"Document limit reached ({tenant.max_documents} documents)"
        )
    
    try:
        # 2. READ FILE CONTENT
        content = await file.read()
        file_extension = file.filename.split('.')[-1].lower()
        
        if file_extension == 'pdf':
            try:
                from pypdf import PdfReader
                import io
                pdf_reader = PdfReader(io.BytesIO(content))
                file_text = ""
                for page in pdf_reader.pages:
                    file_text += page.extract_text() + "\n"
                if not file_text.strip():
                    raise ValueError("PDF appears to be empty or unreadable")
            except ImportError:
                raise HTTPException(status_code=500, detail="PDF support not installed. Run: pip install pypdf")
        
        elif file_extension in ['docx', 'doc']:
            try:
                from docx import Document as DocxDocument
                import io
                docx = DocxDocument(io.BytesIO(content))
                file_text = "\n".join([para.text for para in docx.paragraphs if para.text.strip()])
                if not file_text.strip():
                    raise ValueError("DOCX appears to be empty")
            except ImportError:
                raise HTTPException(status_code=500, detail="DOCX support not installed. Run: pip install python-docx")
        
        elif file_extension == 'txt':
            try:
                file_text = content.decode('utf-8')
            except UnicodeDecodeError:
                try:
                    file_text = content.decode('latin-1')
                except:
                    raise HTTPException(status_code=400, detail="Unable to decode text file. Please ensure it's UTF-8 or Latin-1 encoded.")
        
        else:
            raise HTTPException(status_code=400, detail=f"Unsupported file type: .{file_extension}. Supported types: pdf, docx, txt")
        
        if not file_text or len(file_text.strip()) < 10:
            raise HTTPException(status_code=400, detail="Document is too short or empty. Please upload a document with meaningful content.")
        
        # 3. PARSE TAGS
        tag_list = [tag.strip() for tag in tags.split(",")] if tags else []
        
        # 4. CREATE DATABASE RECORD
        db_document = Document(
            id=f"doc_{uuid.uuid4().hex[:12]}",
            tenant_id=tenant_id,
            filename=file.filename,
            file_size=len(content),
            file_type=file_extension,
            status=DocumentStatus.PROCESSING,
            category=category,
            tags=tag_list,
            description=description
        )
        db.add(db_document)
        await db.commit()
        await db.refresh(db_document)
        
        # 5. PREPARE METADATA FOR QDRANT
        metadata = db_document.to_metadata_dict()
        
        # 6. CREATE LLAMAINDEX DOCUMENT
        llama_doc = LlamaDocument(text=file_text, metadata=metadata, id_=db_document.id)
        
        # 7. SPLIT INTO CHUNKS
        splitter = SentenceSplitter(chunk_size=512, chunk_overlap=50)
        nodes = splitter.get_nodes_from_documents([llama_doc])
        
        for node in nodes:
            node.metadata.update(metadata)
        
        # 8. INDEX IN QDRANT
        vector_store = QdrantVectorStore(
            client=qdrant_client,
            aclient=async_qdrant_client,
            collection_name=settings.QDRANT_COLLECTION
        )
        index = VectorStoreIndex.from_vector_store(vector_store)
        index.insert_nodes(nodes)
        
        # 9. UPDATE DATABASE STATUS
        db_document.status = DocumentStatus.COMPLETED
        db_document.chunk_count = len(nodes)
        db_document.processed_at = datetime.utcnow()
        await db.commit()

        # ← Trigger webhook
        from app.utils.webhooks import trigger_webhook
        await trigger_webhook(
        db=db,
        tenant_id=tenant_id,
        event="document.uploaded",
        payload={
        "document_id": db_document.id,
        "filename": db_document.filename,
        "chunks_created": len(nodes)
    }
)
        
        return DocumentUploadResponse(
            status="success",
            document_id=db_document.id,
            filename=db_document.filename,
            file_size=db_document.file_size,
            chunks_created=len(nodes),
            metadata=metadata
        )
        
    except HTTPException:
        raise
        
    except Exception as e:
        if 'db_document' in locals():
            try:
                db_document.status = DocumentStatus.FAILED
                db_document.error_message = str(e)
                await db.commit()

                from app.utils.webhooks import trigger_webhook
                await trigger_webhook(
                    db=db,
                    tenant_id=tenant_id,
                    event="document.failed",
                    payload={
                    "document_id": db_document.id,
                    "filename": db_document.filename,
                    "error": str(e)
                    }
                )    
            except:
                await db.rollback()
        
        print(f"Upload error for tenant {tenant_id}: {type(e).__name__}: {e}")
        raise HTTPException(status_code=500, detail=f"Document upload failed: {type(e).__name__}")


@router.post("/documents/{document_id}/search")
async def search_in_document(
    document_id: str,
    query: str,
    tenant_id: str = Depends(get_current_tenant_id),
    db: AsyncSession = Depends(get_db),
    qdrant_client = Depends(get_qdrant_client_dependency),
    async_qdrant_client = Depends(get_async_qdrant_client_dependency)
):
    """Search within a specific document only."""
    
    # 1. VERIFY DOCUMENT EXISTS AND BELONGS TO USER
    result = await db.execute(
        select(Document).where(
            Document.id == document_id,
            Document.tenant_id == tenant_id
        )
    )
    document = result.scalar_one_or_none()
    
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
    
    # 2. CHECK DOCUMENT STATUS
    if document.status != DocumentStatus.COMPLETED:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Document not ready. Status: {document.status.value}"
        )
    
    # 3. CREATE REQUEST WITH DOCUMENT FILTER
    request = QueryRequest(
        query=query,
        document_ids=[document_id]
    )
    
    # 4. CALL THE MAIN QUERY FUNCTION
    try:
        result = await ask_question(
            request=request,
            tenant_id=tenant_id,
            qdrant_client=qdrant_client,
            async_qdrant_client=async_qdrant_client
        )
        
        # 5. ENHANCE RESPONSE WITH DOCUMENT INFO
        return {
            "answer": result.answer,
            "document": {
                "id": document.id,
                "filename": document.filename,
                "category": document.category,
                "tags": document.tags,
                "chunk_count": document.chunk_count
            },
            "sources": result.sources,
            "filters_applied": result.filters_applied
        }
    
    except Exception as e:
        print(f"Search error in document {document_id}: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Search failed: {type(e).__name__}"
        )


@router.get("/documents", response_model=DocumentListResponse)
async def list_documents(
    page: int = 1,
    page_size: int = 50,
    category: Optional[str] = None,
    status_filter: Optional[DocumentStatus] = None,
    tenant_id: str = Depends(get_current_tenant_id),
    db: AsyncSession = Depends(get_db)
):
    """List all documents for the current tenant with optional filtering"""
    
    # Build query
    stmt = select(Document).where(Document.tenant_id == tenant_id)
    
    if category:
        stmt = stmt.where(Document.category == category)
    if status_filter:
        stmt = stmt.where(Document.status == status_filter)
    
    # Get total count
    count_stmt = select(func.count()).select_from(Document).where(Document.tenant_id == tenant_id)
    if category:
        count_stmt = count_stmt.where(Document.category == category)
    if status_filter:
        count_stmt = count_stmt.where(Document.status == status_filter)
    
    total_result = await db.execute(count_stmt)
    total = total_result.scalar()
    
    # Get paginated results
    stmt = stmt.offset((page - 1) * page_size).limit(page_size)
    result = await db.execute(stmt)
    documents = result.scalars().all()
    
    return DocumentListResponse(
        total=total,
        documents=documents,
        page=page,
        page_size=page_size
    )


@router.get("/documents/{document_id}", response_model=DocumentResponse)
async def get_document(
    document_id: str,
    tenant_id: str = Depends(get_current_tenant_id),
    db: AsyncSession = Depends(get_db)
):
    """Get details of a specific document"""
    result = await db.execute(
        select(Document).where(
            Document.id == document_id,
            Document.tenant_id == tenant_id
        )
    )
    document = result.scalar_one_or_none()
    
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
    
    return document


@router.patch("/documents/{document_id}", response_model=DocumentResponse)
async def update_document_metadata(
    document_id: str,
    update_data: DocumentUpdateRequest,
    tenant_id: str = Depends(get_current_tenant_id),
    db: AsyncSession = Depends(get_db),
    qdrant_client = Depends(get_qdrant_client_dependency)
):
    """Update document metadata (category, tags, description)"""
    result = await db.execute(
        select(Document).where(
            Document.id == document_id,
            Document.tenant_id == tenant_id
        )
    )
    document = result.scalar_one_or_none()
    
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
    
    # Update fields if provided
    if update_data.category is not None:
        document.category = update_data.category
    if update_data.tags is not None:
        document.tags = update_data.tags
    if update_data.description is not None:
        document.description = update_data.description
    
    await db.commit()
    await db.refresh(document)
    
    # Update metadata in Qdrant
    try:
        from qdrant_client.models import SetPayload, FieldCondition, Filter, MatchValue
        
        updated_metadata = document.to_metadata_dict()
        
        qdrant_client.set_payload(
            collection_name=settings.QDRANT_COLLECTION,
            payload=updated_metadata,
            points_selector=Filter(
                must=[
                    FieldCondition(
                        key="document_id",
                        match=MatchValue(value=document_id)
                    )
                ]
            )
        )
    except Exception as e:
        print(f"Warning: Failed to update Qdrant metadata: {e}")
    
    return document


@router.delete("/documents/{document_id}")
async def delete_document(
    document_id: str,
    tenant_id: str = Depends(get_current_tenant_id),
    db: AsyncSession = Depends(get_db),
    qdrant_client = Depends(get_qdrant_client_dependency)
):
    """Delete a document from both database and vector store"""
    result = await db.execute(
        select(Document).where(
            Document.id == document_id,
            Document.tenant_id == tenant_id
        )
    )
    document = result.scalar_one_or_none()
    
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
    
    try:
        # Delete from Qdrant
        from qdrant_client.models import Filter, FieldCondition, MatchValue
        
        qdrant_client.delete(
            collection_name=settings.QDRANT_COLLECTION,
            points_selector=Filter(
                must=[
                    FieldCondition(
                        key="document_id",
                        match=MatchValue(value=document_id)
                    )
                ]
            )
        )
        
        # Delete from database
        await db.delete(document)
        await db.commit()
        
        return {"status": "success", "message": f"Document {document_id} deleted"}
        
    except Exception as e:
        await db.rollback()
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to delete document: {str(e)}"
        )


@router.post("/ask", response_model=QueryResponse)
async def ask_question(
    request: QueryRequest,
    tenant_id: str = Depends(get_current_tenant_id),
    qdrant_client = Depends(get_qdrant_client_dependency), 
    async_qdrant_client = Depends(get_async_qdrant_client_dependency), 
    db: AsyncSession = Depends(get_db)
):
    """Query documents with optional filtering"""
    if not settings.GEMINI_API_KEY:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, 
            detail="GEMINI_API_KEY not configured"
        )
    
    start_time = datetime.now()
    

    try:
        Settings.llm = Gemini(
            model="gemini-2.0-flash",
            api_key=settings.GEMINI_API_KEY
        )

        vector_store = QdrantVectorStore(
            client=qdrant_client, 
            aclient=async_qdrant_client,
            collection_name=settings.QDRANT_COLLECTION
        )
        index = VectorStoreIndex.from_vector_store(vector_store)

        # Build dynamic filters
        metadata_filters = build_metadata_filters(tenant_id, request)
        
        # Apply filters to retriever
        retriever = index.as_retriever(
            similarity_top_k=5, 
            filters=metadata_filters
        )

        reranker = SentenceTransformerRerank(
            model="cross-encoder/ms-marco-TinyBERT-L-2",
            top_n=3
        )
        
        query_engine = RetrieverQueryEngine(
            retriever=retriever,
            node_postprocessors=[reranker],
        )

        response = await query_engine.aquery(request.query)

        # Extract sources with rich metadata
        sources = []
        for node in response.source_nodes:
            # Convert score to native Python float if it exists
            score = None
            if hasattr(node, 'score') and node.score is not None:
                score = float(node.score)  # Convert numpy.float32 to Python float
            
            source_info = {
                "document_id": node.metadata.get("document_id", "N/A"),
                "file_name": node.metadata.get("file_name", "N/A"),
                "category": node.metadata.get("category", "uncategorized"),
                "tags": node.metadata.get("tags", []),
                "content_preview": node.text[:200] + "...",
                "score": score
            }
            sources.append(source_info)
        answer = str(response)

        end_time = datetime.now()
        response_time = int((end_time - start_time).total_seconds() * 1000)

        query_log = QueryLog(
            tenant_id=tenant_id,
            query_text=request.query,
            filters_applied=get_applied_filters_summary(request),
            answer_length=len(answer),
            sources_count=len(sources),
            response_time_ms=response_time,
            success=True
        )
        db.add(query_log)
        await db.commit()

        return QueryResponse(
            answer=str(response),
            sources=sources,
            filters_applied=get_applied_filters_summary(request)
        )

    except Exception as e:

        end_time = datetime.utcnow()
        response_time = int((end_time - start_time).total_seconds() * 1000)
        query_log = QueryLog(
            tenant_id=tenant_id,
            query_text=request.query,
            filters_applied=get_applied_filters_summary(request),
            answer_length=len(answer),
            sources_count=len(sources),
            response_time_ms=response_time,
            success=False,
            error_message=str(e)
        )
        db.add(query_log)
        await db.commit()

        print(f"Query Error for Tenant {tenant_id}: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Query processing failed: {type(e).__name__}"
        )